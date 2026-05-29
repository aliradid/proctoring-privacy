"""Build the manuscript as a Word document conforming to the Informatica
(Slovenia) author template (manuscript/informatica_template.docx).

The template defines the journal's named paragraph styles (title,
Normal-no indent, abstract, abstract Si, heading 1, heading 2,
Bibliography). We open the template as a document, clear its body, and emit
the content of `manuscript/manuscript.md` using those styles directly so the
result respects the journal's typography out of the box. A STYLE_MAP
translates the logical element names used below onto the template's actual
style names; section headings and bibliography entries are auto-numbered by
the styles themselves.

Run: `python3 code/build_informatica.py`
Output: `manuscript/Article_Informatica.docx`
"""
from __future__ import annotations

import copy
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, Twips

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = ROOT / "manuscript" / "informatica_template.docx"
MS_PATH = ROOT / "manuscript" / "manuscript.md"
OUT = ROOT / "manuscript" / "Article_Informatica.docx"
FIG_DIR = ROOT / "figures"

# Running-head metadata. The template ships with placeholder header text
# ("Enter short title...", volume 23/37, year 1999/2013, author "D. Torkar")
# wired to TITLE/AUTHOR/DOCPROPERTY fields. We overwrite the placeholders with
# this paper's values so the page headers render correctly (see _patch_parts).
SHORT_TITLE = "Privacy-Preserving Multi-Modal Proctoring for Online Exams"
AUTHOR_RUNNING = "R. Ali et al."
JOURNAL_LINE = "Informatica 50 (2026) xxx–yyy"  # volume/year current; pages set by editor


def _patch_parts(data: dict):
    """Patch the template's package parts (headers, doc properties, settings)
    so the running heads show THIS paper instead of the template placeholders.

    Pages 2+ use header1 (even) and header2 (odd), which contain a static
    journal line plus PAGE/TITLE/AUTHOR fields. The first page uses header3,
    which the template builds from DOCPROPERTY fields; we replace it with a
    clean static line to avoid depending on field recomputation for the volume,
    year and page range.
    """
    def get(k):
        return data[k].decode("utf-8")

    def put(k, s):
        data[k] = s.encode("utf-8")

    # header1 (even pages): journal line + AUTHOR field (cached "Borut").
    if "word/header1.xml" in data:
        h1 = get("word/header1.xml")
        h1 = h1.replace("Informatica 23 (1999) xxx–yyy", JOURNAL_LINE)
        h1 = h1.replace("<w:t>Borut</w:t>", "<w:t>%s</w:t>" % AUTHOR_RUNNING)
        put("word/header1.xml", h1)

    # header2 (odd pages): TITLE field (cached placeholder) + journal line.
    if "word/header2.xml" in data:
        h2 = get("word/header2.xml")
        h2 = h2.replace(
            "<w:t>Enter short title in File/Properties/Summary</w:t>",
            "<w:t>%s</w:t>" % SHORT_TITLE,
        )
        h2 = h2.replace("<w:t>23</w:t>", "<w:t>50</w:t>")
        h2 = h2.replace(
            '<w:t xml:space="preserve"> (1999) xxx–yyy</w:t>',
            '<w:t xml:space="preserve"> (2026) xxx–yyy</w:t>',
        )
        put("word/header2.xml", h2)

    # header3 (first page): replace the DOCPROPERTY-driven paragraph with a
    # clean static journal line + the current page number.
    if "word/header3.xml" in data:
        h3 = get("word/header3.xml")
        new_p = (
            '<w:p><w:pPr><w:pStyle w:val="Header"/>'
            '<w:tabs><w:tab w:val="right" w:pos="9524"/></w:tabs></w:pPr>'
            '<w:r><w:rPr><w:sz w:val="18"/></w:rPr><w:tab/>'
            '<w:t xml:space="preserve">%s  </w:t></w:r>'
            '<w:r><w:rPr><w:rStyle w:val="PageNumber"/><w:b/></w:rPr>'
            '<w:fldChar w:fldCharType="begin"/></w:r>'
            '<w:r><w:rPr><w:rStyle w:val="PageNumber"/><w:b/></w:rPr>'
            '<w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'
            '<w:r><w:rPr><w:rStyle w:val="PageNumber"/><w:b/></w:rPr>'
            '<w:fldChar w:fldCharType="separate"/></w:r>'
            '<w:r><w:rPr><w:rStyle w:val="PageNumber"/><w:b/><w:noProof/></w:rPr>'
            '<w:t>1</w:t></w:r>'
            '<w:r><w:rPr><w:rStyle w:val="PageNumber"/><w:b/></w:rPr>'
            '<w:fldChar w:fldCharType="end"/></w:r></w:p>'
        ) % JOURNAL_LINE
        h3 = re.sub(r"<w:p\b.*?</w:p>", new_p, h3, count=1, flags=re.S)
        put("word/header3.xml", h3)

    # core.xml: drive the TITLE/AUTHOR fields and strip the template owner trace.
    if "docProps/core.xml" in data:
        core = get("docProps/core.xml")
        core = core.replace(
            "Enter short title in File/Properties/Summary", SHORT_TITLE
        )
        core = core.replace("D. Torkar et al.", AUTHOR_RUNNING)
        core = core.replace("Drago Torkar", "Radid Ali")
        put("docProps/core.xml", core)

    # custom.xml: tidy the volume/year/author traces (used by unreferenced
    # headers only, but kept consistent so document metadata is clean).
    if "docProps/custom.xml" in data:
        cust = get("docProps/custom.xml")
        cust = cust.replace("<vt:i4>37</vt:i4>", "<vt:i4>50</vt:i4>")
        cust = cust.replace("<vt:i4>2013</vt:i4>", "<vt:i4>2026</vt:i4>")
        cust = cust.replace("T.J. Author et al.", AUTHOR_RUNNING)
        cust = cust.replace("Enter short title into File/Properties", SHORT_TITLE)
        put("docProps/custom.xml", cust)

    # Note: we deliberately do NOT set <w:updateFields/>, which would make Word
    # raise a modal "update fields?" prompt on open (it blocks PDF automation).
    # The TITLE/AUTHOR/journal-line results are patched directly above, and PAGE
    # fields recompute automatically during pagination, so no refresh is needed.

# Map the builder's logical style names onto the ACTUAL style display-names
# defined in the official Informatica template (word_example). python-docx
# looks styles up by display name, so these are the w:name values, not ids.
STYLE_MAP = {
    "I_Title": "title",
    "I_Authors": "Normal-no indent",
    "I_Text": "Normal-no indent",
    "I_BodyIndent": "Normal",
    "I_Keywords": "Normal-no indent",
    "I_Received": "Normal-no indent",
    "I_Abstract": "abstract",
    "I_Abstract_SI": "abstract Si",
    "I_SectionTitle": "heading 1",
    "I_SubSectionTitle": "heading 2",
    "I_FigureCaption": "Normal-no indent",
    "I_TableCaption": "Normal-no indent",
    "I_Figure": "Normal-no indent",
    "I_Table": "Normal",
    "I_References": "heading 1",
    "I_Bibliography": "Bibliography",
}

# Same figure insertion plan as before, but with Informatica-style captions
# Single-column width (≈ 3.15"). Wide content that does not fit will be wrapped
# in a one-column section break (see `_open_one_column` / `_open_two_columns`).
FIG_WIDTH_INCH = 3.15
WIDE_FIG_WIDTH_INCH = 6.4
# Tables that have these many or more columns are emitted in a one-column
# section so they can use the full page width.
WIDE_TABLE_COL_THRESHOLD = 4


def _new_section_break(doc, num_cols: int, header_refs=None, footer_refs=None,
                       title_pg: bool = False):
    """Append a continuous section break that switches the column count.

    Word renders this as the document continuing on the same page but with a
    new column setting. When ``header_refs``/``footer_refs``/``title_pg`` are
    supplied, the section also carries those header/footer references and the
    title-page flag, so the running heads from the template are preserved (the
    OOXML schema requires header/footer references before <w:type> and titlePg
    after <w:cols>).
    """
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sect_pr = OxmlElement("w:sectPr")
    # header/footer references must precede <w:type> in the schema sequence
    for ref in (header_refs or []):
        sect_pr.append(copy.deepcopy(ref))
    for ref in (footer_refs or []):
        sect_pr.append(copy.deepcopy(ref))
    # continuous section break
    type_el = OxmlElement("w:type")
    type_el.set(qn("w:val"), "continuous")
    sect_pr.append(type_el)
    pgsz = OxmlElement("w:pgSz")
    pgsz.set(qn("w:w"), "11906")
    pgsz.set(qn("w:h"), "16838")
    sect_pr.append(pgsz)
    pgmar = OxmlElement("w:pgMar")
    pgmar.set(qn("w:top"), "1644")
    pgmar.set(qn("w:right"), "1077")
    pgmar.set(qn("w:bottom"), "1134")
    pgmar.set(qn("w:left"), "1304")
    pgmar.set(qn("w:header"), "936")
    pgmar.set(qn("w:footer"), "720")
    sect_pr.append(pgmar)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(num_cols))
    cols.set(qn("w:space"), "284")
    sect_pr.append(cols)
    if title_pg:
        sect_pr.append(OxmlElement("w:titlePg"))
    pPr.append(sect_pr)


# Figures are placed by explicit [[FIG:key]] markers in the markdown, which is
# robust to prose edits. Each marker is replaced (in document order) by the
# image + an auto-numbered caption. Captions are keyed here.
FIGURE_CAPTIONS = {
    "architecture": "End-to-end privacy-preserving proctoring pipeline.",
    "si_timeline": "Simulated 60-second session illustrating the trajectory of the visual, acoustic and behavioural risk components and the resulting Suspicion Index.",
    "visual_pr_curves": "Precision-recall curves of YOLOv8-s on the COCO val2017 proctoring subset (217 images, 1,164 annotations).",
    "visual_tp_fp": "Per-class true-positive, false-positive and false-negative counts after greedy IoU 0.5 matching.",
    "acoustic_roc_pr": "ROC and PR curves for the Whisper-Base content-free acoustic module on the LibriSpeech-derived benchmark.",
    "acoustic_per_class": "Per-class distribution of the secondary-speaker probability on the 400-clip LibriSpeech benchmark.",
    "behaviour_distributions": "Per-scenario distribution of the geometric descriptors (yaw amplitude and lip variance) computed from MediaPipe FaceMesh landmarks of real Pexels face anchors.",
    "fusion_comparison": "Cross-validated F1 of the five fusion strategies on real-modality features (error bars = standard deviation over 5 folds).",
    "fusion_roc": "Pooled 5-fold cross-validation ROC curves of all fusion strategies on real-modality features.",
    "fusion_confusion": "Confusion matrix of the random-forest fusion model under 5-fold cross-validation on the 300 real-feature sessions.",
    "fusion_ablation": "Ablation of the random-forest fusion model: leave-one-modality-out (upper) and single-modality (lower).",
    "latency": "Per-module inference latency on the Apple M4 Max host (Mac Studio, 14-core CPU, 32-core integrated GPU, 36 GB unified memory) with both MPS GPU acceleration and CPU-only configurations.",
    "robustness": "Robustness study: class-presence recall versus gamma and scale for the visual subsystem, and F1 plus score distributions versus SNR for the acoustic subsystem under real LibriSpeech noise injection.",
    "privacy": "Adversarial privacy verification on 30 real LibriSpeech speakers: re-identification AUC achieved by a raw-audio attacker (proxy of conventional proctoring) versus a metadata-only attacker (the proposed system).",
}
# Figures that must span the full page width (rendered in a 1-column section).
WIDE_FIGURE_KEYS = {
    "architecture", "si_timeline", "acoustic_roc_pr", "fusion_roc",
    "latency", "robustness",
}
FIG_MARKER_RE = re.compile(r"^\s*\[\[FIG:([a-z_]+)\]\]\s*$")


def _apply_style(p, style_name: str, doc):
    """Apply a template style, translating the builder's logical name to the
    real Informatica template style name; fall back to Normal if missing."""
    target = STYLE_MAP.get(style_name, style_name)
    try:
        p.style = doc.styles[target]
    except KeyError:
        try:
            p.style = doc.styles["Normal"]
        except KeyError:
            pass


def _add_run(p, text: str, bold: bool = False, italic: bool = False):
    r = p.add_run(text)
    if bold:
        r.font.bold = True
    if italic:
        r.font.italic = True
    return r


# --- LaTeX-to-Word conversion --------------------------------------------
# Word cannot render LaTeX, so inline math ($...$) and display math ($$...$$)
# from the markdown must be converted to Unicode symbols plus real
# subscript/superscript runs, otherwise the reader sees raw "$\mathrm{...}$".
_LATEX_SYM = {
    r"\theta": "θ", r"\phi": "φ", r"\varphi": "φ", r"\sigma": "σ", r"\alpha": "α",
    r"\beta": "β", r"\gamma": "γ", r"\delta": "δ", r"\mu": "μ", r"\lambda": "λ",
    r"\epsilon": "ε", r"\varepsilon": "ε", r"\rho": "ρ", r"\tau": "τ",
    r"\in": "∈", r"\notin": "∉", r"\geq": "≥", r"\leq": "≤", r"\ge": "≥",
    r"\le": "≤", r"\times": "×", r"\cdot": "·", r"\approx": "≈", r"\pm": "±",
    r"\cup": "∪", r"\cap": "∩", r"\to": "→", r"\rightarrow": "→", r"\infty": "∞",
    r"\leftarrow": "←", r"\neq": "≠", r"\sum": "Σ", r"\forall": "∀", r"\exists": "∃",
}


def _latex_preprocess(s: str) -> str:
    """Turn a LaTeX math string into Unicode, keeping `_`/`^` for scripts."""
    # scripts whose argument is \text{...} / \mathrm{...} -> brace group
    s = re.sub(r"([_^])\\(?:text|mathrm|mathbf|operatorname)\{([^}]*)\}", r"\1{\2}", s)
    # scripts whose argument is a bare operator word -> brace group
    # (use a letter-lookahead, not \b: '_' is a word char so \b fails before it)
    s = re.sub(r"([_^])\\(max|min|sec|dur|off|ovl|whisp|sup|inf)(?![a-zA-Z])", r"\1{\2}", s)
    # \bar over a following symbol -> symbol + combining overline
    s = re.sub(r"\\bar\\([a-zA-Z]+)",
               lambda m: _LATEX_SYM.get("\\" + m.group(1), m.group(1)) + "̅", s)
    s = re.sub(r"\\bar\{([^}]*)\}", lambda m: m.group(1) + "̅", s)
    s = re.sub(r"\\hat\{([^}]*)\}", lambda m: m.group(1) + "̂", s)
    # font wrappers -> inner content
    for cmd in ("mathrm", "mathbf", "mathsf", "mathit", "text", "operatorname", "boldsymbol"):
        s = re.sub(r"\\" + cmd + r"\{([^}]*)\}", r"\1", s)
    # blackboard / calligraphic
    s = s.replace(r"\mathbb{R}", "ℝ").replace(r"\mathbb{N}", "ℕ").replace(r"\mathbb{Z}", "ℤ")
    s = re.sub(r"\\mathbb\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\mathcal\{([^}]*)\}", r"\1", s)
    # standalone operator words (letter-lookahead so '\max_' / '\max}' match)
    for w in ("argmax", "argmin", "max", "min", "log", "exp", "sin", "cos", "tanh"):
        s = re.sub(r"\\" + w + r"(?![a-zA-Z])", w, s)
    # delimiters / spacing commands
    s = re.sub(r"\\(bigg?|Bigg?|left|right)\b", "", s)
    s = s.replace(r"\{", "{").replace(r"\}", "}")
    s = s.replace(r"\,", " ").replace(r"\;", " ").replace(r"\:", " ").replace(r"\!", "").replace(r"\ ", " ")
    # symbol map (longest first to avoid partial hits)
    for k in sorted(_LATEX_SYM, key=len, reverse=True):
        s = s.replace(k, _LATEX_SYM[k])
    return s


def _emit_latex(p, latex: str):
    """Render a LaTeX math string into runs on paragraph p, with real
    subscript/superscript. A `{` is treated as a script group only when it
    directly follows `_` or `^`; otherwise it is a literal brace."""
    s = _latex_preprocess(latex)
    buf = ""
    i = 0

    def flush():
        nonlocal buf
        if buf:
            _add_run(p, buf)
            buf = ""

    while i < len(s):
        ch = s[i]
        if ch in "_^" and i + 1 < len(s):
            flush()
            is_sup = ch == "^"
            j = i + 1
            if s[j] == "{":
                k = s.find("}", j)
                if k == -1:
                    content, i = s[j + 1:], len(s)
                else:
                    content, i = s[j + 1:k], k + 1
            else:
                content, i = s[j], j + 1
            r = p.add_run(content)
            if is_sup:
                r.font.superscript = True
            else:
                r.font.subscript = True
        else:
            buf += ch
            i += 1
    flush()


def _emit_rich(p, text: str):
    """Emit text into paragraph p, handling inline math ($...$) and markdown
    bold/italic. Inline math is converted via _emit_latex."""
    for part in re.split(r"(\$[^$]+\$)", text):
        if not part:
            continue
        if len(part) >= 2 and part.startswith("$") and part.endswith("$"):
            _emit_latex(p, part[1:-1])
        else:
            for seg in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", part):
                if not seg:
                    continue
                if seg.startswith("**") and seg.endswith("**"):
                    _add_run(p, seg[2:-2], bold=True)
                elif seg.startswith("*") and seg.endswith("*"):
                    _add_run(p, seg[1:-1], italic=True)
                else:
                    _add_run(p, seg)


def _add_text_paragraph(doc, text: str, style: str = "I_Text"):
    p = doc.add_paragraph()
    _apply_style(p, style, doc)
    _emit_rich(p, text)
    return p


def _add_figure(doc, key: str, fig_no: int):
    fig_name = f"fig_{key}.png"
    fig_path = FIG_DIR / fig_name
    caption_text = FIGURE_CAPTIONS.get(key, "")
    if not fig_path.exists():
        print(f"  WARNING missing figure {fig_name}")
        return
    wide = key in WIDE_FIGURE_KEYS
    if wide:
        # Close the surrounding 2-column body section.
        _new_section_break(doc, 2)
    p = doc.add_paragraph()
    _apply_style(p, "I_Figure", doc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    width = Inches(WIDE_FIG_WIDTH_INCH if wide else FIG_WIDTH_INCH)
    p.add_run().add_picture(str(fig_path), width=width)
    cap = doc.add_paragraph()
    _apply_style(cap, "I_FigureCaption", doc)
    _add_run(cap, f"Figure {fig_no}: ", bold=True)
    _add_run(cap, caption_text)
    if wide:
        # Close the wide 1-column section that contained this figure.
        _new_section_break(doc, 1)


# Full text width of a one-column section and of one body column (twips).
FULL_TEXT_TWIPS = 9525
BODY_COL_TWIPS = 4620


def _fit_table(table, header, rows, total_twips: int):
    """Force a fixed layout and distribute column widths to fill the width.

    Without this, python-docx tables default to an auto/autofit layout and Word
    shrinks them to content, crushing columns until words break mid-character.
    Single-character columns get a small fixed width; the rest are sized by the
    longer of their header and (capped) content length, then scaled to fill
    total_twips.
    """
    def clean(x):
        return str(x).replace("**", "").strip()

    ncols = len(header)
    maxlen = []
    for j in range(ncols):
        cells = [header[j]] + [row[j] if j < len(row) else "" for row in rows]
        maxlen.append(max((len(clean(c)) for c in cells), default=1))
    NARROW = 300  # ~0.2 inch, enough for a single letter or a tick
    narrow = [j for j in range(ncols) if maxlen[j] <= 2]
    wide = [j for j in range(ncols) if maxlen[j] > 2]
    widths = [0] * ncols
    rem = total_twips - NARROW * len(narrow)
    if wide and rem > 0:
        def weight(j):
            return max(len(clean(header[j])), min(maxlen[j], 22))
        wsum = sum(weight(j) for j in wide)
        for j in narrow:
            widths[j] = NARROW
        for j in wide:
            widths[j] = int(rem * weight(j) / wsum)
    else:
        widths = [total_twips // ncols] * ncols
    # Fixed layout (python-docx places <w:tblLayout w:type="fixed"/> correctly).
    table.autofit = False
    table.allow_autofit = False
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for col, wd in zip(grid.findall(qn("w:gridCol")), widths):
            col.set(qn("w:w"), str(wd))
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            if j < len(widths):
                cell.width = Twips(widths[j])


def _add_table_md(doc, header, rows, table_no: int, title: str):
    wide = len(header) >= WIDE_TABLE_COL_THRESHOLD
    if wide:
        _new_section_break(doc, 2)
    # Caption above
    cap = doc.add_paragraph()
    _apply_style(cap, "I_TableCaption", doc)
    _add_run(cap, f"Table {table_no}: ", bold=True)
    _add_run(cap, title)
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    try:
        table.style = doc.styles["Table Grid"]
    except KeyError:
        # Apply manual borders if Table Grid isn't shipped with the template
        from docx.oxml import OxmlElement as _OE
        tbl_pr = table._tbl.tblPr
        borders = _OE("w:tblBorders")
        for tag in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = _OE(f"w:{tag}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "4")
            b.set(qn("w:color"), "000000")
            borders.append(b)
        tbl_pr.append(borders)
    for j, h in enumerate(header):
        c = table.rows[0].cells[j]
        c.text = ""
        para = c.paragraphs[0]
        _apply_style(para, "I_Table", doc)
        _add_run(para, h.strip(), bold=True)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            c = table.rows[i].cells[j]
            c.text = ""
            para = c.paragraphs[0]
            _apply_style(para, "I_Table", doc)
            text = val.strip()
            is_bold = text.startswith("**") and text.endswith("**")
            if is_bold:
                text = text[2:-2]
            _add_run(para, text, bold=is_bold)
    # Repeat the header row on continuation pages and stop individual rows from
    # splitting across a page break, so a long table that straddles a page
    # boundary keeps its column labels instead of orphaning bare numeric rows.
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(tbl_header)
    for r in table.rows:
        r._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    # Fill the available width with sensible per-column widths so columns do not
    # collapse and wrap mid-word.
    _fit_table(table, header, rows, FULL_TEXT_TWIPS if wide else BODY_COL_TWIPS)
    if wide:
        _new_section_break(doc, 1)


def _parse_table(lines, idx):
    header_line = lines[idx]
    sep = lines[idx + 1]
    if not re.match(r"\s*\|.*\|\s*$", sep):
        return None
    header = [c.strip() for c in header_line.strip().strip("|").split("|")]
    rows = []
    i = idx + 2
    while i < len(lines) and re.match(r"\s*\|.*\|\s*$", lines[i]):
        rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
        i += 1
    return header, rows, i


def render():
    work_path = OUT.with_suffix(".docx")
    # 1) Convert .dotx -> .docx by patching the content-type only.
    with zipfile.ZipFile(TEMPLATE, "r") as zin:
        data = {name: zin.read(name) for name in zin.namelist()}
    ct = data["[Content_Types].xml"].decode()
    ct = ct.replace(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
    )
    data["[Content_Types].xml"] = ct.encode()
    # Rewrite the running-head placeholders / properties to this paper.
    _patch_parts(data)
    tmp = ROOT / "manuscript" / "_informatica_seed.docx"
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, payload in data.items():
            zout.writestr(name, payload)

    # 2) Open in python-docx and remove every pre-existing body element so we
    #    keep the styles, page setup and section properties but start with a
    #    clean canvas.
    doc = Document(tmp)
    body = doc.element.body
    # Capture the template's first (title-page) section: its titlePg flag and
    # its even/default/first header references. Clearing the body below deletes
    # the paragraph that carried them, so we restore them onto the title section
    # we emit ourselves; every later section then inherits them (link-to-previous).
    title_hdr_refs, title_ftr_refs, title_pg = [], [], False
    for sp in body.iter(qn("w:sectPr")):
        if sp.find(qn("w:headerReference")) is not None or sp.find(qn("w:titlePg")) is not None:
            title_hdr_refs = [copy.deepcopy(r) for r in sp.findall(qn("w:headerReference"))]
            title_ftr_refs = [copy.deepcopy(r) for r in sp.findall(qn("w:footerReference"))]
            title_pg = sp.find(qn("w:titlePg")) is not None
            break
    # Preserve the sectPr at the very end of the body
    sect_pr = body.find(qn("w:sectPr"))
    # Drop its lone header/footer references so the final section inherits the
    # title section's running heads instead of the template's example header.
    if sect_pr is not None:
        for ref in sect_pr.findall(qn("w:headerReference")) + sect_pr.findall(qn("w:footerReference")):
            sect_pr.remove(ref)
    for child in list(body):
        if child is sect_pr:
            continue
        body.remove(child)

    md = MS_PATH.read_text().splitlines()
    fig_no = 1
    table_no = 1
    i = 0
    in_authors = False
    title_done = False
    references_started = False
    body_two_columns_started = False
    # The template's body typography distinguishes paragraphs by a first-line
    # indent (style "Normal"), except the first paragraph after a heading, which
    # is flush-left (style "Normal-no indent"). Track that here.
    first_body_para = False

    while i < len(md):
        line = md[i]
        s = line.rstrip()

        # H1 — title
        if s.startswith("# "):
            p = doc.add_paragraph()
            _apply_style(p, "I_Title", doc)
            _add_run(p, s[2:].strip(), bold=True)
            title_done = True
            i += 1
            continue

        # Author block
        if s.startswith("**Radid Ali**"):
            p = doc.add_paragraph()
            _apply_style(p, "I_Authors", doc)
            _add_run(p, "Radid Ali, Ghazouani Mohamed and El habib Benlahmar")
            aff = doc.add_paragraph()
            _apply_style(aff, "I_Text", doc)
            _add_run(aff, "Department of Mathematics and Computer Science, Hassan II University, Faculty of Sciences Ben M'sik, Casablanca, Morocco", italic=True)
            mail = doc.add_paragraph()
            _apply_style(mail, "I_Text", doc)
            _add_run(mail, "E-mail: ali.radid-etu@etu.univh2c.ma, ghazouani.fsbm@gmail.com, h.benlahmer@gmail.com", italic=True)
            corr = doc.add_paragraph()
            _apply_style(corr, "I_Text", doc)
            _add_run(corr, "Corresponding author: Radid Ali (ali.radid-etu@etu.univh2c.ma)", italic=True)
            # Per the Informatica template, a blank line separates the author
            # block from the Keywords/Received block, the Keywords paragraph
            # carries 12 pt of space after it, and another blank line separates
            # Received from the abstract.
            _apply_style(doc.add_paragraph(), "I_Text", doc)  # blank separator
            kw_text = ""
            for ln in md:
                if ln.strip().startswith("**Keywords**"):
                    kw_text = ln.strip()[len("**Keywords**"):].lstrip(" :-")
                    break
            kw = doc.add_paragraph()
            _apply_style(kw, "I_Keywords", doc)
            kw.paragraph_format.space_after = Pt(12)
            _add_run(kw, "Keywords: ", bold=True)
            _add_run(kw, kw_text)
            rec = doc.add_paragraph()
            _apply_style(rec, "I_Received", doc)
            _add_run(rec, "Received: ")
            _apply_style(doc.add_paragraph(), "I_Text", doc)  # blank separator before abstract
            # skip until the `## Abstract` heading (the abstract is emitted next)
            j = i + 1
            while j < len(md) and not md[j].strip().startswith("## Abstract"):
                j += 1
            i = j
            continue

        # Abstract handling: H2 "## Abstract" then a body paragraph; we render
        # the body as an I_Abstract paragraph (no separate heading line).
        if s.startswith("## Abstract"):
            # Consume the abstract body (the next non-empty line) as I_Abstract
            j = i + 1
            while j < len(md) and md[j].strip() == "":
                j += 1
            abstract_lines = []
            while j < len(md) and md[j].strip() and not md[j].strip().startswith(("##", "**Keywords**")):
                abstract_lines.append(md[j].strip())
                j += 1
            abstract_text = " ".join(abstract_lines)
            p = doc.add_paragraph()
            _apply_style(p, "I_Abstract", doc)
            _emit_rich(p, abstract_text)
            i = j
            continue
        if s.strip().startswith("**Keywords**"):
            # Keywords were already emitted in the front-matter block (before
            # the abstract, per template). Here we emit only the Slovenian
            # Povzetek, which follows the English abstract. The translation was
            # prepared with machine-translation help and should be checked by a
            # native Slovenian speaker before final acceptance.
            ps = doc.add_paragraph()
            _apply_style(ps, "I_Abstract_SI", doc)
            _add_run(ps, "Povzetek: ", bold=True)
            _add_run(ps, (
                "Hitra uporaba spletnih platform za ocenjevanje je pospešila uvajanje "
                "proctoring sistemov, ki temeljijo na umetni inteligenci. Večina komercialnih "
                "rešitev se zanaša na vsiljivo biometrično obdelavo, kot so neprekinjeno prepoznavanje "
                "obrazov, izločanje glasovnih odtisov in prepis govora, kar je v nasprotju s "
                "strogimi predpisi o varstvu podatkov, kot sta evropska Splošna uredba o varstvu "
                "podatkov (GDPR) in maroška uredba CNDP, zakon 09-08, ki tovrstne podatke "
                "uvrščata med občutljive osebne podatke. V članku predstavljamo večmodalno "
                "ogrodje za nadzor izpitov, ki ohranja zasebnost in ne shranjuje biometričnih "
                "predlog, surovega videa ali jezikovne vsebine. Sistem povezuje YOLOv8-s za "
                "kontekstualno zaznavanje predmetov, MediaPipe FaceMesh za neidentificirajoče "
                "vedenjske namige (odklon pogleda, drža glave, gibanje ustnic) in Whisper-Base "
                "kodirnik v načinu brez prepisa govora. Vsi izhodi modulov so pretvorjeni v "
                "standardiziran tok metapodatkov in obdelani z modelom za izračun indeksa suma, "
                "katerega strategija je izbrana empirično med petimi različicami. Na realni "
                "podmnožici slik COCO val2017 (217 slik, 1.164 oznak) YOLOv8-s doseže makro F1 "
                "0,606 in povprečni AP@0,5 0,519. Pri 400 zvočnih posnetkih, sestavljenih iz "
                "korpusa LibriSpeech, kodirnik Whisper-Base v brezvsebinskem načinu doseže ROC "
                "AUC 0,774 za prisotnost dodatnega govorca (F1 = 0,78) in AUC 0,935 za "
                "zaznavanje šepetanja. Pri 5-kratnem prečnem preverjanju na 300 sejah s "
                "preverjenimi modalnostnimi izhodi se naključni gozd uvrsti kot najboljša "
                "strategija s F1 = 0,905 ± 0,041 (95-odstotni interval [0,876, 0,931]); njegova "
                "prednost pred preprostim pravilom tehtane vsote (F1 = 0,899) ni statistično "
                "značilna (McNemar p = 0,37), je pa značilno boljši od različic MLP in naivnega "
                "Bayesa. Na napravi Apple M4 Max (Mac Studio) cevovod deluje pri "
                "23,7 sličicah na sekundo s pospeševanjem MPS in 19,4 sličicah na sekundo v "
                "načinu samo CPU. Sistem porabi največ 1.102 MB pomnilnika. Pri preverjanju "
                "zasebnosti z 30 resničnimi govorci napadalec s surovim zvokom doseže AUC 0,994 "
                "za ponovno identifikacijo, medtem ko enak napadalec, omejen na tok "
                "metapodatkov, doseže le 0,478 (pod naključno mejo). To kvantitativno potrjuje "
                "lastnost ogrodja, da ne shranjuje biometričnih podatkov, in ohrani skladnost z "
                "9. členom GDPR ter z zakonom 09-08 maroške CNDP."
            ))
            # End of title block: insert a continuous section break that makes
            # everything *before* it appear in one column. The final sectPr
            # at the end of the document keeps the body in two columns. This
            # break also restores the template's title-page headers (titlePg +
            # first/default/even references) that clearing the body removed.
            _new_section_break(doc, 1, header_refs=title_hdr_refs,
                               footer_refs=title_ftr_refs, title_pg=title_pg)
            body_two_columns_started = True
            i += 1
            continue

        # Section / subsection headings
        if s.startswith("## ") and "References" not in s and "Acknowledgements" not in s and "Data and Code" not in s:
            heading = s[3:].strip()
            # The Informatica I_SectionTitle / I_SubSectionTitle styles
            # auto-number sections, so we strip the leading "N. " or "N.M "
            # so we don't render "7 7. Discussion".
            heading = re.sub(r"^\d+(?:\.\d+)?\.?\s+", "", heading)
            p = doc.add_paragraph()
            _apply_style(p, "I_SectionTitle", doc)
            _add_run(p, heading, bold=True)
            first_body_para = True
            i += 1
            continue
        if s.startswith("## References"):
            p = doc.add_paragraph()
            # I_References is the template's UN-numbered back-matter heading
            # style; using I_SectionTitle here would auto-number it ("12").
            _apply_style(p, "I_References", doc)
            _add_run(p, "References", bold=True)
            references_started = True
            i += 1
            continue
        if s.startswith("## Acknowledgements") or s.startswith("## Data and Code Availability"):
            p = doc.add_paragraph()
            _apply_style(p, "I_References", doc)
            _add_run(p, s[3:].strip(), bold=True)
            first_body_para = True
            i += 1
            continue
        if s.startswith("### "):
            heading = s[4:].strip()
            heading = re.sub(r"^\d+(?:\.\d+)?\.?\s+", "", heading)
            p = doc.add_paragraph()
            _apply_style(p, "I_SubSectionTitle", doc)
            _add_run(p, heading, bold=True)
            first_body_para = True
            i += 1
            continue

        # Explicit figure marker: [[FIG:key]]
        fig_marker = FIG_MARKER_RE.match(line)
        if fig_marker:
            _add_figure(doc, fig_marker.group(1), fig_no)
            fig_no += 1
            i += 1
            continue

        # Tables
        if re.match(r"\s*\|.*\|\s*$", line):
            parsed = _parse_table(md, i)
            if parsed:
                header, rows, next_i = parsed
                # Look back for a "Table N. ..." caption line above
                title = ""
                for back in range(i - 1, max(0, i - 4), -1):
                    if md[back].strip().startswith("**Table"):
                        title = re.sub(r"^\*\*Table \d+\.\s*", "", md[back].strip()).rstrip("*").strip()
                        break
                _add_table_md(doc, header, rows, table_no, title)
                table_no += 1
                i = next_i
                continue

        # Equation blocks (single-line $$...$$) -> centered, math-rendered
        if s.startswith("$$") and s.endswith("$$") and len(s) > 4:
            p = doc.add_paragraph()
            _apply_style(p, "I_Text", doc)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _emit_latex(p, s[2:-2].strip())
            i += 1
            continue
        if s.startswith("$$"):
            j = i + 1
            buf = []
            while j < len(md) and not md[j].strip().startswith("$$"):
                buf.append(md[j])
                j += 1
            p = doc.add_paragraph()
            _apply_style(p, "I_Text", doc)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _emit_latex(p, " ".join(x.strip() for x in buf).strip())
            i = j + 1
            continue

        # Bulleted / numbered lists
        if s.lstrip().startswith("- "):
            p = doc.add_paragraph()
            _apply_style(p, "I_Text", doc)
            _add_run(p, "• ")
            _emit_rich(p, s.lstrip()[2:])
            i += 1
            continue
        m_num = re.match(r"^(\d+)\.\s+(.*)", s.lstrip())
        if m_num:
            p = doc.add_paragraph()
            _apply_style(p, "I_Text", doc)
            _add_run(p, f"{m_num.group(1)}. ", bold=True)
            _emit_rich(p, m_num.group(2))
            i += 1
            continue

        # References. The template's Bibliography style auto-numbers each entry
        # as "[1] ", "[2] ", ... (numId 7, lvlText "[%1] "), so we strip the
        # literal "[N]" from the markdown and let the style supply it. Entries
        # are in citation order, so the auto-numbers match the in-text [N].
        if references_started and re.match(r"^\[\d+\]", s.strip()):
            entry = re.sub(r"^\[\d+\]\s*", "", s.strip())
            p = doc.add_paragraph()
            _apply_style(p, "I_Bibliography", doc)
            _emit_rich(p, entry)
            i += 1
            continue

        # Skip blank lines + table-caption-only lines (already handled)
        if s.strip() == "":
            i += 1
            continue
        if s.strip().startswith("**Table"):
            i += 1
            continue

        # Collapse soft-wrapped paragraph
        para_lines = [s]
        j = i + 1
        while j < len(md):
            nxt = md[j]
            if not nxt.strip() or nxt.lstrip().startswith(("#", "|", "- ", "$$", "```")) or FIG_MARKER_RE.match(nxt) or re.match(r"^\d+\.\s", nxt.strip()) or (references_started and re.match(r"^\[\d+\]", nxt.strip())):
                break
            para_lines.append(nxt)
            j += 1
        para_text = " ".join(p.strip() for p in para_lines)
        # First paragraph after a heading is flush-left; the rest are indented,
        # matching the template's body typography.
        _add_text_paragraph(doc, para_text, style=("I_Text" if first_body_para else "I_BodyIndent"))
        first_body_para = False
        i = j

    doc.save(work_path)
    print(f"Wrote {work_path}")
    print(f"  ({fig_no - 1} figures, {table_no - 1} tables)")


if __name__ == "__main__":
    render()
