"""Render cover_letter.md as a plain Word document on Times New Roman 11pt."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "manuscript" / "cover_letter.md"
OUT = ROOT / "manuscript" / "Cover_Letter_Informatica.docx"


def set_font(run, size=11, bold=False, italic=False, font="Times New Roman"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)


def add_text(doc, text, **kw):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for part in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text):
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        italic = part.startswith("*") and part.endswith("*") and not bold
        run = p.add_run(part.strip("*"))
        set_font(run, bold=bold, italic=italic, **kw)
    return p


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.0)
        section.left_margin = section.right_margin = Cm(2.5)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)

    lines = SRC.read_text().splitlines()
    i = 0
    while i < len(lines):
        ln = lines[i].rstrip()
        if ln.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(ln[2:].strip())
            set_font(r, size=13, bold=True)
            i += 1
            continue
        if ln.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(ln[3:].strip())
            set_font(r, size=12, bold=True)
            i += 1
            continue
        if ln.strip().startswith("- "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            txt = "• " + ln.strip()[2:]
            for part in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", txt):
                if not part:
                    continue
                bold = part.startswith("**") and part.endswith("**")
                italic = part.startswith("*") and part.endswith("*") and not bold
                r = p.add_run(part.strip("*"))
                set_font(r, bold=bold, italic=italic)
            i += 1
            continue
        m_num = re.match(r"^(\d+)\.\s+(.*)", ln.strip())
        if m_num:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            r = p.add_run(f"{m_num.group(1)}. ")
            set_font(r, bold=True)
            for part in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", m_num.group(2)):
                if not part:
                    continue
                bold = part.startswith("**") and part.endswith("**")
                italic = part.startswith("*") and part.endswith("*") and not bold
                r = p.add_run(part.strip("*"))
                set_font(r, bold=bold, italic=italic)
            i += 1
            continue
        if ln.strip() == "":
            i += 1
            continue
        # Collapse soft-wrapped paragraphs
        buf = [ln]
        j = i + 1
        while j < len(lines) and lines[j].strip() and not lines[j].lstrip().startswith(("#", "- ")) and not re.match(r"^\d+\.\s", lines[j].strip()):
            buf.append(lines[j])
            j += 1
        add_text(doc, " ".join(b.strip() for b in buf))
        i = j
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
